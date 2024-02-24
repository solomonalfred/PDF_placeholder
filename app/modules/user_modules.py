import shutil
from constants.variables import *
from constants.api_items import *
from fastapi import UploadFile, File
from docx import Document
import re
import PyPDF2
import pandas as pd
from io import BytesIO
from pathlib import Path
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))
from core.core_object import Core


def save_file(username: str, input_file_data: UploadFile = File(...)):
    try:
        filename = input_file_data.filename
        path = os.path.join(FILE_FOLDER, username, filename)
        with open(f'{path}', "wb") as buffer:
            shutil.copyfileobj(input_file_data.file, buffer)
        return path
    except:
        return ''


def get_tags(file_path: str):
    tags = []
    tb = []
    pattern = r"<<(.*?)>>"

    def process(doc):
        for p in doc.paragraphs:
            inline = p.runs
            flag = False
            placeholder_tags(inline, flag)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process(cell)

    def placeholder_tags(inline, flag):
        reg = ""
        for i in range(len(inline)):
            left = inline[i].text.find("<<")
            left_part = inline[i].text.find("<")
            right = inline[i].text.find(">>")
            right_part = inline[i].text.find(">")
            if left != -1 and right != -1: reg = inline[i].text
            elif left != -1:
                reg += inline[i].text
                flag = True
                continue
            elif right != -1 and flag:
                reg += inline[i].text
                flag = False
            elif left == -1 and right == -1 and flag:
                reg += inline[i].text
                continue
            elif left_part != -1:
                if len(inline) > i + 1 and inline[i + 1].text[0] == "<":
                    reg += inline[i].text
                    flag = True
                    continue
                else: continue
            elif right_part != -1 and flag:
                if len(inline) > i + 1 and inline[i + 1].text[0] == ">":
                    reg += inline[i].text
                    continue
                elif reg[-1] == ">": reg += inline[i].text
                else:
                    reg += inline[i].text
                    continue
            else:
                continue

            replace_tags(reg)
            reg = ""

    def replace_tags(reg):
        matches = re.findall(pattern, reg)
        for match in matches:
            if match.find('_tb') != -1:
                tb.append(match)
            else:
                tags.append(match)

    try:
        doc = Document(file_path)
        process(doc)
        for section in doc.sections:
            header = section.header
            process(header)
            footer = section.footer
            process(footer)
    except:
        pass
    return [list(set(tags)), list(set(tb))]


def keys_response(res):
    data_tags = dict_tags(res[Details.RESPONCE][0])
    data_tb = dict_tags(res[Details.RESPONCE][1])
    result = {Details.COUNT_TAGS: len(data_tags.keys()),
              Details.COUNT_TB: len(data_tb.keys()),
              Details.KEYS: data_tags,
              Details.TB: data_tb}
    return result


def error_response_render(res, newfilename):
    if res[Details.RESPONCE] == msg.INSUFFICIENT_FUNDS:
        return [402, {msg.MSG: msg.INSUFFICIENT_FUNDS}]
    if res[Details.RESPONCE] == msg.TP_DELETED:
        return [400, {msg.MSG: msg.TEMPLATE_NOT_EXISTS}]
    newfilename = newfilename.replace('.pdf', '').replace('.docx', '')
    if res[Details.RESPONCE].find(newfilename) == -1:
        return [500, {msg.MSG: msg.INTERNAL_SERVER_ERROR}]
    return [200, res[Details.RESPONCE]]


def temp_excel_builder(res):
    df = pd.DataFrame(res[Details.TRANSACTIONS])
    output = BytesIO()
    with pd.ExcelWriter(output, engine=Excel_items.ENGINE) as writer:
        df.to_excel(writer, sheet_name=Excel_items.SHEET, index=False)
        workbook = writer.book
        worksheet = writer.sheets[Excel_items.SHEET]
        header_format = workbook.add_format({Excel_items.HEADER: True})
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
        writer._save()
    output.seek(0)
    return output


def balancer_process_response(file_path, newfilename, filename, username, data):
    if file_path is None or file_path[Table_items.DELETED]:
        return [False, 400, {Details.RESPONCE: msg.TP_DELETED}]
    file_path = file_path[Table_items.PATH]
    if len(newfilename) == 0:
        newfilename = filename
    newfilename = newfilename.replace('.pdf', '').replace('.docx', '')
    filler = Core(username, file_path, newfilename, data).process()
    if filler.find(newfilename) == -1:
        return [False, 400, {Details.RESPONCE: msg.INTERNAL_SERVER_ERROR}]
    else:
        file = Path(filler).name
        file_size = os.path.getsize(filler)
        count = count_pages(filler)
        return [True, 200, {Details.RESPONCE: [file, file_size, count, filler]}]


def balancer_transaction_fill_record(tmp):
    t = dict()
    t[Table_items.TYPE] = tmp[Table_items.TYPE]
    t[Table_items.BALANCE] = tmp[Table_items.BALANCE]
    if tmp[Table_items.UNLIMITED]:
        t[Table_items.AMOUNT] = Table_items.UNLIMITED
    else:
        t[Table_items.AMOUNT] = tmp[Table_items.AMOUNT]
    if t[Table_items.TYPE] == Details.CREDIT:
        t[Table_items.FILE] = tmp[Table_items.FILE]
        t[Table_items.PAGE_PROCESSED] = tmp[Table_items.PAGE_PROCESSED]
    else:
        t[Table_items.FILE] = "-"
        t[Table_items.TEMPLATE] = "-"
        t[Table_items.PAGE_PROCESSED] = "-"
    t[Table_items.CRATED_AT] = tmp[Table_items.CRATED_AT]
    return t


def dict_tags(tags):
    tag = {}
    for i in tags: tag[i] = ""
    return tag


def transform_user(old: dict, new: dict):
    old.update(new)
    return old


def delete_transform(old: dict, new: dict):
    if len(old) == 1 and old.get(new['key']) is not None:
        return {}
    if len(old) == 1 and old.get(new['key']) is None:
        return old
    del old[new['key']]
    return old


def count_pages(path: str):
    pdf_file = open(path, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    num_pages = len(pdf_reader.pages)
    pdf_file.close()
    return num_pages
