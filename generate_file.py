from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import random
import string

def generate_random_string(min_length=1, max_length=8):
    length = random.randint(min_length, max_length)
    random_string = ''.join(random.choices(string.ascii_letters + string.digits, k=length))
    return f"<<{random_string}>>"

def create_word_document(num_pages, filename):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    for _ in range(num_pages):
        for _ in range(13):
            paragraph = doc.add_paragraph()
            while len(paragraph.text) < 70:
                paragraph.add_run(generate_random_string() + " ")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

    if doc.paragraphs[-1].text == '':
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    doc.save(filename)


def create_docx(elements_per_page, number_of_pages):
    doc = Document()

    for page in range(number_of_pages):
        for element in range(elements_per_page):
            # Генерируем случайный текст для элемента (от 2 до 20 символов)
            random_text = ''.join(random.choices(string.ascii_letters + string.digits, k=random.randint(2, 20)))
            # Добавляем элемент (текст) в документ, обрамленный << и >>
            doc.add_paragraph(f'<<{random_text}>>')

        # Добавляем разрыв страницы после каждой страницы, кроме последней
        if page < number_of_pages - 1:
            doc.add_section()

    # Сохраняем документ
    doc.save('output.docx')


if __name__ == "__main__":
    # create_word_document(100, 'random_text_document.docx')
    create_docx(25, 1)
