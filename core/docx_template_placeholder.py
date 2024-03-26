import os.path
import re
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from typing import Dict

from docx.shared import Pt

from constants.variables import *
from constants.msg import ErrorType
from constants.core_items import *
from core.pdf_convertor import Convert2PDF
from core.error_block import ErrorBlocker


class DocxTemplatePlaceholder:

    def __init__(self,
                 username: str,
                 template: str,
                 newfilename: str,
                 tags: Dict
                 ):
        self.error = ErrorType.ok
        try:
            self.template_document = Document(template)
            self.file_name = template.split('/')[-1]
            self.replace_tags = self.__prepare_tags(tags)
            self.username = username
            self.new = newfilename
        except:
            self.error = ErrorType.incorrect_doc

    def process(self):
        if self.error:
            return ErrorBlocker().process(self.error)
        try:
            self.__process(self.template_document, self.replace_tags)
            self.__placeholder_footer_header()
            tmp_filename = self.file_name.replace(Process_items.DOCX, Process_items.NEW_DOCX)
            path = os.path.join(FILE_FOLDER, self.username, tmp_filename)
            self.template_document.save(path)
            return Convert2PDF(path, self.new, self.username).DocxToPdf()
        except:
            self.error = ErrorType.internal_error
            return ErrorBlocker().process(self.error)

    def __process(self, doc, tags, tables=True):

        for p in doc.paragraphs:
            inline = p.runs
            flag = False
            if tables:
                table_flag = self.__generate_table(tags[1], p, doc)
                if table_flag: continue
            self.__placeholder_tags(inline, tags[0], flag)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.__process(cell, tags)

    def __generate_table(self, tb, p, doc):
        for regex, replace in tb.items():
            if regex.search(p.text):
                p.text = ''
                table = doc.add_table(rows=1, cols=len(replace[0]))
                self.__fill_table(table, replace)
                self.__insert_table_before_paragraph(table, p)
                return True
        return False

    def __fill_table(self, table, data):
        table.style = Table_items.TABLE_GRID

        hdr_cells = table.rows[0].cells
        keys = list(data[0].keys())
        for i, key in enumerate(keys):
            hdr_cells[i].text = key.capitalize().replace('_', ' ')
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            self.__set_cell_borders(hdr_cells[i])
        for item in data:
            row_cells = table.add_row().cells
            for idx, key in enumerate(keys):
                row_cells[idx].text = str(item[key])
                self.__set_cell_borders(row_cells[idx])

        self.__format_table_cells(table)
        self.__set_column_widths(table)

    def __set_column_widths(self, table):
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(1)

    def __format_table_cells(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def __insert_table_before_paragraph(self, table, paragraph):
        tbl_element = table._tbl
        p_element = paragraph._p
        p_parent = p_element.getparent()
        p_parent.insert(p_parent.index(p_element), tbl_element)

    def __set_cell_borders(self, cell, border_sz=4):
        sides = [Table_items.TOP,
                 Table_items.LEFT,
                 Table_items.BOTTOM,
                 Table_items.RIGHT]
        for side in sides:
            border_elm = OxmlElement(Table_items.OXML.format(side))
            border_att = {
                f'{qn(Table_items.VALUE)}': Table_items.SINGLE,
                f'{qn(Table_items.SZ)}': str(border_sz),
                f'{qn(Table_items.SPACE)}': Table_items.ZERO,
                f'{qn(Table_items.COLOR)}': Table_items.AUTO,
            }
            border_elm.attrib.update(border_att)
            cell._tc.get_or_add_tcPr().append(border_elm)

    def __placeholder_tags(self, inline, tags, flag):
        reg = ""
        for i in range(len(inline)):
            borders = self.__check_regex_borders(inline[i].text)
            if borders[0] != -1 and borders[2] != -1:
                if len(reg):
                    inline[i - 1].text = reg
                reg = inline[i].text
            elif borders[0] != -1:
                reg = self.__splicing_tag(inline, i, reg)
                flag = True
                continue
            elif borders[2] != -1 and flag:
                reg += inline[i].text
                flag = False
            elif borders[0] == -1 and borders[2] == -1 and flag and borders[1] == -1 and borders[3] == -1:
                reg = self.__splicing_tag(inline, i, reg)
                continue
            elif borders[1] != -1:
                future_part = len(inline) > i + 1 and inline[i + 1].text[0] == Inline.PART_LEFT_BORDER
                past_part = len(reg) and reg[-1] == Inline.PART_LEFT_BORDER
                if future_part or past_part:
                    reg = self.__splicing_tag(inline, i, reg)
                    flag = True
                continue
            elif borders[3] != -1 and flag:
                if len(inline) > i + 1 and inline[i + 1].text[0] == Inline.PART_RIGHT_BORDER:
                    reg = self.__splicing_tag(inline, i, reg)
                    continue
                elif reg[-1] == Inline.PART_RIGHT_BORDER:
                    reg += inline[i].text
                else:
                    reg = self.__splicing_tag(inline, i, reg)
                    continue
            else:
                continue
            self.__replace_tags(tags, inline[i], reg)
            reg = ""
        if len(reg):
            inline[-1].text = reg


    def __splicing_tag(self, inline, idx,  reg):
        reg += inline[idx].text
        inline[idx].text = ""
        return reg


    def __check_regex_borders(self, text: str):
        left = text.find(Inline.LEFT_BORDER)
        left_part = text.find(Inline.PART_LEFT_BORDER)
        right = text.find(Inline.RIGHT_BORDER)
        right_part = text.find(Inline.PART_RIGHT_BORDER)
        return [left, left_part, right, right_part]

    def __replace_tags(self, tags, inline, reg):
        skip = True
        for regex, replace in tags.items():
            if regex.search(reg):
                text = regex.sub(replace, reg)
                inline.text = text
                skip = False
                break
        if skip:
            inline.text = reg

    def __placeholder_footer_header(self):
        for section in self.template_document.sections:
            self.__process(section.header, self.replace_tags, False)
            self.__process(section.footer, self.replace_tags, False)

    def __prepare_tags(self, tags):
        done_tags = dict()
        done_tb = dict()
        if len(tags.items()) == 0:
            return [done_tags, done_tb]
        done_tags = self.__fill_dict_placeholders(Process_items.KEYS, tags)
        done_tb = self.__fill_dict_placeholders(Process_items.TABLES, tags)
        return [done_tags, done_tb]


    def __regex_dict(self, name: str, tags: dict):
        placeholders = dict()
        for regex, replace in tags[name].items():
            placeholders[re.compile(Process_items.REGEX_FORM.format(regex))] = replace
        return placeholders


    def __fill_dict_placeholders(self, name: str, tags: dict):
        done_dict = dict()
        check_flag = tags.get(name)
        if check_flag is not None:
            done_dict = self.__regex_dict(name, tags)
        return done_dict
